from telethon.sync import TelegramClient
import socks
from docx import Document
import re
import yagmail
import schedule
import time
from collections import Counter

# ——— 配置区 ———
api_id = 29504380
api_hash = '6af8088819d2de076e5712d799129131'
channel = 'dtao_txn'
EMAIL = '470167795@qq.com'
EMAIL_PASSWORD = 'nnopdolnsfsgbhac'  # 请确保这是正确的 QQ SMTP 授权码

# 本地的 SOCKS5 代理设置，请替换为你当前的 IP 和端口
PROXY = (socks.SOCKS5, '127.0.0.1', 7890, True)

# 仅抓取最新 100 条消息
MSG_LIMIT = 100
# 只发送 1-3 条异常 Swap 记录
ANOMALY_MIN = 1
ANOMALY_MAX = 3
# Swap 匹配样式（示例：Swap 12.34 TAO for 56.78 α）
SWAP_PATTERN = r"Swap ([\d\.]+) TAO for ([\d\.]+) α"

# ——— 功能函数定义 ———

def fetch_messages():
    print("🔌 正在通过代理连接 Telegram…")
    with TelegramClient('tao_session', api_id, api_hash, proxy=PROXY) as client:
        me = client.get_me()
        print("✅ 已连接账号：", me.username or me.first_name)
        messages = client.iter_messages(channel, limit=MSG_LIMIT)
        return [msg.text or "" for msg in messages]


def normalize_swap_message(msg: str) -> str:
    """对 Swap 消息做轻度规范化，避免同一内容的轻微空格差异导致计数偏差。"""
    return re.sub(r"\s+", " ", msg.strip())


def extract_unique_swaps(messages):
    """从最新消息中提取仅出现一次的 Swap 记录（按完整消息文本计数）。"""
    swap_messages = []
    for msg in messages:
        if re.search(SWAP_PATTERN, msg):
            swap_messages.append(normalize_swap_message(msg))

    counts = Counter(swap_messages)
    unique_swaps = [m for m, c in counts.items() if c == 1]

    # 按消息在原始顺序出现的先后进行保序过滤（保留原始列表顺序）
    ordered_unique = []
    seen = set()
    for m in swap_messages:
        if m in counts and counts[m] == 1 and m not in seen:
            ordered_unique.append(m)
            seen.add(m)

    # 只取 1-3 条之间（若不足 1 条则返回空列表，由调度逻辑决定是否发送）
    if len(ordered_unique) == 0:
        return []
    return ordered_unique[:ANOMALY_MAX]


def make_doc(unique_swaps):
    doc = Document()
    doc.add_heading('异常 Swap 交易（最新 100 条，仅出现一次）', level=1)
    for swap in unique_swaps:
        doc.add_paragraph(swap)
    path = 'tao_swap_anomalies.docx'
    doc.save(path)
    return path


def send_mail(path, unique_swaps):
    try:
        yag = yagmail.SMTP(
            user=EMAIL,
            password=EMAIL_PASSWORD,
            host='smtp.qq.com',
            port=465,
            smtp_ssl=True
        )
        body_lines = [
            '以下为在频道 dtao_txn 最新 100 条消息中仅出现一次的 Swap 记录：',
            ''
        ] + unique_swaps
        yag.send(
            to=EMAIL,
            subject='[TAO] 异常 Swap 交易告警',
            contents='\n'.join(body_lines),
            attachments=path
        )
        print('✅ 邮件已发送')
    except Exception as e:
        print(f"❌ 邮件发送失败: {e}")


def job():
    print("🛠 任务开始执行…")
    messages = fetch_messages()
    unique_swaps = extract_unique_swaps(messages)
    # 仅当满足 1-3 条之间才发送（0 条不发，多于 3 条只取前 3 条）
    if ANOMALY_MIN <= len(unique_swaps) <= ANOMALY_MAX:
        path = make_doc(unique_swaps)
        send_mail(path, unique_swaps)
    else:
        print("⚠️ 未发现 1-3 条的异常 Swap（0 条或超过上限），等待下一次调度…")


# ——— 调度逻辑 ———

# 设置为每 10 分钟执行一次
schedule.every(10).minutes.do(job)

print("✅ Bot 已启动，等待调度循环…")
while True:
    schedule.run_pending()
    time.sleep(5)